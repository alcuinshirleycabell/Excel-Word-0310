import os
import re
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from difflib import SequenceMatcher

import pandas as pd
from docx import Document


SIGNATURE_TEXT = "@徐静妍18616309259"


def find_best_match(word_text: str, excel_cols: set, threshold=0.6):
    """
    智能匹配Word中的占位符到Excel列名。
    处理简写vs全称的情况，如 "日均总资产" -> "日均总资产(近一年)"
    """
    if word_text in excel_cols:
        return word_text, 1.0

    candidates = []

    for col in excel_cols:
        ratio = SequenceMatcher(None, word_text, col).ratio()

        # 规则1：完全包含关系
        if word_text in col and len(word_text) >= 4:
            ratio = 0.95
        elif col in word_text and len(col) >= 4:
            ratio = 0.9

        # 规则2：去掉括号后匹配
        col_clean = re.sub(r'\([^)]*\)', '', col).strip()
        word_clean = re.sub(r'\([^)]*\)', '', word_text).strip()

        if word_clean == col_clean:
            ratio = 0.98
        elif word_clean in col_clean or col_clean in word_clean:
            ratio = max(ratio, 0.88)

        if ratio >= threshold:
            candidates.append((col, ratio, len(col)))

    if not candidates:
        return None, 0

    # 优先选择相似度高的，相似度相同时选择长度更接近的
    candidates.sort(key=lambda x: (-x[1], -abs(len(x[0]) - len(word_text))))
    return candidates[0][0], candidates[0][1]


def build_mapping(row: pd.Series, excel_cols: set = None) -> dict:
    """
    构建占位符映射。
    支持三种格式：纯文本、__字段__、{字段}
    """
    mapping = {}
    for col, val in row.items():
        col = col.strip()
        v = "" if pd.isna(val) else str(val)
        mapping[col] = v
        mapping[f"__{col}__"] = v
        mapping["{" + col + "}"] = v

    return mapping


def replace_smart(paragraph, mapping: dict, excel_cols: set = None):
    """
    智能替换段落中的占位符。
    支持简写形式自动匹配完整列名。
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return 0

    new_text = full_text

    # 步骤1：预处理，将简写占位符替换为完整列名
    if excel_cols:
        # 提取可能的中文占位符（包括括号）
        potential = re.findall(r'[\u4e00-\u9fa5]+(?:[（(][^）)]+[）)])?(?:[\u4e00-\u9fa5]+)?', new_text)
        potential = sorted(set(potential), key=len, reverse=True)

        for ph in potential:
            if ph not in excel_cols:
                match, ratio = find_best_match(ph, excel_cols)
                if match and match != ph and ratio >= 0.9:
                    new_text = new_text.replace(ph, match)

    # 步骤2：使用mapping替换（按key长度降序，避免部分匹配）
    for key in sorted(mapping.keys(), key=len, reverse=True):
        if key in new_text:
            new_text = new_text.replace(key, mapping[key])

    if new_text == full_text:
        return 0

    # 写回段落
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)

    return 1


def _iter_paragraphs_in_cell(cell):
    """递归遍历单元格中的所有段落"""
    for p in cell.paragraphs:
        yield p
    for table in cell.tables:
        for row in table.rows:
            for inner_cell in row.cells:
                yield from _iter_paragraphs_in_cell(inner_cell)


def iter_all_paragraphs(doc: Document):
    """遍历文档中的所有段落（包括表格、页眉页脚）"""
    for p in doc.paragraphs:
        yield p

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs_in_cell(cell)

    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from _iter_paragraphs_in_cell(cell)
        for p in section.footer.paragraphs:
            yield p
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from _iter_paragraphs_in_cell(cell)


def replace_in_doc(doc: Document, mapping: dict, excel_cols: set = None):
    """替换文档中所有占位符"""
    total = 0
    for paragraph in iter_all_paragraphs(doc):
        total += replace_smart(paragraph, mapping, excel_cols)
    return total


class App:
    def __init__(self, root: tk.Tk):
        self.root = root
        root.title("批量Word生成工具")
        root.geometry("900x560")

        self.template_path = tk.StringVar()
        self.excel_path = tk.StringVar()
        self.out_dir = tk.StringVar()
        self.status = tk.StringVar(value="准备就绪")
        self.filename_col = tk.StringVar(value="")
        self.filename_pattern = tk.StringVar(value="")

        frm = ttk.Frame(root, padding=12)
        frm.pack(fill=tk.BOTH, expand=True)

        ttk.Label(frm, text="Word模板(.docx)：").grid(row=0, column=0, sticky=tk.W)
        ttk.Entry(frm, textvariable=self.template_path, width=74).grid(row=0, column=1, padx=6)
        ttk.Button(frm, text="选择模板", command=self.pick_template).grid(row=0, column=2)

        ttk.Label(frm, text="Excel数据(.xlsx)：").grid(row=1, column=0, sticky=tk.W, pady=(6, 0))
        ttk.Entry(frm, textvariable=self.excel_path, width=74).grid(row=1, column=1, padx=6, pady=(6, 0))
        ttk.Button(frm, text="选择Excel", command=self.pick_excel).grid(row=1, column=2, pady=(6, 0))

        ttk.Label(frm, text="输出目录：").grid(row=2, column=0, sticky=tk.W, pady=(6, 0))
        ttk.Entry(frm, textvariable=self.out_dir, width=74).grid(row=2, column=1, padx=6, pady=(6, 0))
        ttk.Button(frm, text="选择目录", command=self.pick_outdir).grid(row=2, column=2, pady=(6, 0))

        ttk.Label(frm, text="按哪一列命名（可选）：").grid(row=3, column=0, sticky=tk.W, pady=(10, 0))
        ttk.Entry(frm, textvariable=self.filename_col, width=24).grid(
            row=3, column=1, sticky=tk.W, padx=6, pady=(10, 0)
        )

        ttk.Label(frm, text="自定义文件名模板（可选）：").grid(row=4, column=0, sticky=tk.W, pady=(6, 0))
        ttk.Entry(frm, textvariable=self.filename_pattern, width=74).grid(
            row=4, column=1, padx=6, pady=(6, 0), sticky=tk.W
        )
        ttk.Label(
            frm,
            text="示例：证书_{姓名}_{编号}；可用 {序号} 表示行号（从1开始）",
            foreground="#666666",
        ).grid(row=5, column=1, sticky=tk.W)

        ttk.Label(frm, text="Excel字段预览：").grid(row=6, column=0, sticky=tk.W, pady=(12, 0))
        self.listbox = tk.Listbox(frm, height=10, width=78)
        self.listbox.grid(row=7, column=0, columnspan=2, sticky=tk.W, pady=6)

        btns = ttk.Frame(frm)
        btns.grid(row=8, column=0, columnspan=3, pady=10, sticky=tk.W)
        ttk.Button(btns, text="刷新字段", command=self.load_fields).pack(side=tk.LEFT, padx=6)
        ttk.Button(btns, text="开始生成", command=self.start).pack(side=tk.LEFT, padx=6)

        ttk.Label(frm, textvariable=self.status).grid(row=9, column=0, columnspan=3, sticky=tk.W, pady=(10, 0))

        sign = ttk.Label(root, text=SIGNATURE_TEXT, foreground="#999999")
        sign.place(relx=1.0, rely=1.0, x=-12, y=-8, anchor="se")

    def _ui(self, callback, *args, **kwargs):
        self.root.after(0, lambda: callback(*args, **kwargs))

    def _set_status(self, text: str):
        self._ui(self.status.set, text)

    def pick_template(self):
        path = filedialog.askopenfilename(filetypes=[("Word", "*.docx")])
        if path:
            self.template_path.set(path)
            self.status.set(f"已选择模板：{os.path.basename(path)}")

    def pick_excel(self):
        path = filedialog.askopenfilename(filetypes=[("Excel", "*.xlsx;*.xls")])
        if path:
            self.excel_path.set(path)
            self.status.set(f"已选择Excel：{os.path.basename(path)}")
            self.load_fields()

    def pick_outdir(self):
        directory = filedialog.askdirectory()
        if directory:
            self.out_dir.set(directory)

    def load_fields(self):
        self.listbox.delete(0, tk.END)
        excel_path = self.excel_path.get()
        if not excel_path or not os.path.exists(excel_path):
            self.status.set("请先选择Excel")
            return
        try:
            df = pd.read_excel(excel_path, dtype=str)
            df.columns = [col.strip() for col in df.columns]
            for col in df.columns.tolist():
                self.listbox.insert(tk.END, col)
            self.status.set(f"Excel字段读取成功：{len(df.columns)}列，{len(df)}行数据")
        except Exception as err:
            messagebox.showerror("错误", f"读取Excel失败：{err}")

    @staticmethod
    def sanitize_filename(name: str) -> str:
        name = re.sub(r'[\\/:*?"<>|]+', "_", name).strip()
        return name or "output"

    def render_filename(self, row: pd.Series, idx: int, name_col: str, pattern: str) -> str:
        if pattern:
            rendered = pattern
            for col, val in row.items():
                col = col.strip()
                rendered = rendered.replace("{" + col + "}", str(val or ""))
            rendered = rendered.replace("{序号}", str(idx + 1))
            return self.sanitize_filename(rendered)

        if name_col:
            name_col = name_col.strip()
            return self.sanitize_filename(str(row.get(name_col, "")).strip() or f"output_{idx + 1}")

        return f"output_{idx + 1}"

    def start(self):
        template_path = self.template_path.get()
        excel_path = self.excel_path.get()
        output_dir = self.out_dir.get()

        if not template_path or not os.path.exists(template_path):
            messagebox.showwarning("提示", "请选择Word模板(.docx)")
            return
        if not excel_path or not os.path.exists(excel_path):
            messagebox.showwarning("提示", "请选择Excel数据")
            return
        if not output_dir:
            messagebox.showwarning("提示", "请选择输出目录")
            return

        threading.Thread(target=self.worker, daemon=True).start()

    def worker(self):
        try:
            self._set_status("开始生成...")
            df = pd.read_excel(self.excel_path.get(), dtype=str).fillna("")
            df.columns = [col.strip() for col in df.columns]
            excel_cols = set(df.columns)

            total = len(df)
            if total == 0:
                self._ui(messagebox.showinfo, "提示", "Excel没有数据行")
                self._set_status("无数据")
                return

            name_col = self.filename_col.get().strip()
            pattern = self.filename_pattern.get().strip()

            if name_col and name_col not in df.columns:
                self._set_status(f"文件名列"{name_col}"不存在，改用默认命名")
                name_col = ""

            out_dir = self.out_dir.get()
            os.makedirs(out_dir, exist_ok=True)

            for idx, row in df.iterrows():
                mapping = build_mapping(row, excel_cols)

                doc = Document(self.template_path.get())
                replace_count = replace_in_doc(doc, mapping, excel_cols)

                base = self.render_filename(row, idx, name_col, pattern)
                out_path = os.path.join(out_dir, f"{base}.docx")

                if os.path.exists(out_path):
                    out_path = os.path.join(out_dir, f"{base}_{idx + 1}.docx")

                doc.save(out_path)
                self._set_status(f"已生成 {idx + 1}/{total}：{os.path.basename(out_path)} (替换{replace_count}处)")

            self._ui(messagebox.showinfo, "完成", f"全部生成完成，共 {total} 个文件。")
            self._set_status("生成完成")
        except Exception as err:
            self._ui(messagebox.showerror, "错误", str(err))
            self._set_status("出错：" + str(err))


def run_cli():
    """命令行模式（用于GitHub Actions）"""
    import argparse

    parser = argparse.ArgumentParser(description="Excel批量生成Word工具")
    parser.add_argument("--cli", action="store_true", help="使用命令行模式")
    parser.add_argument("--template", required=True, help="Word模板文件路径")
    parser.add_argument("--excel", required=True, help="Excel数据文件路径")
    parser.add_argument("--output", required=True, help="输出目录")
    parser.add_argument("--filename-col", default="", help="用于文件命名的列名")
    parser.add_argument("--filename-pattern", default="", help="文件名模板")
    args = parser.parse_args()

    print(f"📖 读取Excel: {args.excel}")
    df = pd.read_excel(args.excel, dtype=str).fillna("")
    df.columns = [col.strip() for col in df.columns]
    excel_cols = set(df.columns)

    total = len(df)
    print(f"📊 共 {total} 行数据，{len(df.columns)} 列")
    print(f"📋 列名: {', '.join(list(df.columns)[:5])}{'...' if len(df.columns) > 5 else ''}")

    os.makedirs(args.output, exist_ok=True)

    for idx, row in df.iterrows():
        mapping = build_mapping(row, excel_cols)

        doc = Document(args.template)
        replace_count = replace_in_doc(doc, mapping, excel_cols)

        # 确定文件名
        if args.filename_pattern:
            base = args.filename_pattern
            for col, val in row.items():
                col = col.strip()
                base = base.replace("{" + col + "}", str(val or ""))
            base = base.replace("{序号}", str(idx + 1))
            base = re.sub(r'[\\/:*?"<>|]+', "_", base).strip() or f"output_{idx + 1}"
        elif args.filename_col and args.filename_col.strip() in df.columns:
            base = str(row.get(args.filename_col.strip(), "")).strip() or f"output_{idx + 1}"
            base = re.sub(r'[\\/:*?"<>|]+', "_", base).strip()
        else:
            base = f"output_{idx + 1}"

        out_path = os.path.join(args.output, f"{base}.docx")

        if os.path.exists(out_path):
            out_path = os.path.join(args.output, f"{base}_{idx + 1}.docx")

        doc.save(out_path)
        print(f"✅ [{idx + 1}/{total}] {os.path.basename(out_path)} (替换{replace_count}处)")

    print(f"\n🎉 全部完成！共生成 {total} 个文件")


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1 and sys.argv[1] == "--cli":
        run_cli()
    else:
        root = tk.Tk()
        App(root)
        root.mainloop()
